import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
import json
import subprocess
import os
import shutil
from datetime import datetime

class ResumeGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Resume Generator")
        self.root.geometry("800x600")
        
        # Configure styles
        self.style = ttk.Style()
        self.style.configure('TFrame', background='#f0f0f0')
        self.style.configure('TLabel', background='#f0f0f0', font=('Arial', 10))
        self.style.configure('TButton', font=('Arial', 10))
        self.style.configure('TEntry', font=('Arial', 10))
        self.style.configure('TText', font=('Arial', 10))
        
        self.create_widgets()
        
    def create_widgets(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # File name section
        file_frame = ttk.LabelFrame(main_frame, text="Output File", padding="10")
        file_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(file_frame, text="File Name:").grid(row=0, column=0, sticky=tk.W)
        self.filename_var = tk.StringVar(value="Yallaiah_Data_Engineer_c.docx")
        self.filename_entry = ttk.Entry(file_frame, textvariable=self.filename_var, width=50)
        self.filename_entry.grid(row=0, column=1, padx=5, sticky=tk.W)
        
        ttk.Button(file_frame, text="Browse Location", command=self.browse_location).grid(row=0, column=2, padx=5)
        
        # JSON input section
        json_frame = ttk.LabelFrame(main_frame, text="Resume JSON", padding="10")
        json_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.json_text = tk.Text(json_frame, width=80, height=20, wrap=tk.WORD, font=('Arial', 10))
        self.json_text.pack(fill=tk.BOTH, expand=True)
        
        # Buttons frame
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(buttons_frame, text="Load JSON File", command=self.load_json_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Generate Resume", command=self.generate_resume).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Clear", command=self.clear_fields).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Exit", command=self.root.quit).pack(side=tk.RIGHT, padx=5)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN)
        self.status_bar.pack(fill=tk.X, pady=(5,0))
        
    def browse_location(self):
        initial_dir = "D:/Resumes_Data_Engineers/New_Resumes"
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            title="Save Resume As",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
            initialfile=self.filename_var.get()
        )
        if file_path:
            self.filename_var.set(file_path)
    
    def load_json_file(self):
        file_path = filedialog.askopenfilename(
            title="Select JSON Resume File",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r') as f:
                    json_data = f.read()
                    self.json_text.delete(1.0, tk.END)
                    self.json_text.insert(tk.END, json_data)
                    self.status_var.set(f"Loaded JSON file: {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load JSON file: {str(e)}")
    
    def clear_fields(self):
        self.json_text.delete(1.0, tk.END)
        self.filename_var.set("Yallaiah_Data_Engineer_c.docx")
        self.status_var.set("Fields cleared")
    
    def generate_resume(self):
        json_string = self.json_text.get(1.0, tk.END).strip()
        output_filename = self.filename_var.get().strip()
        
        if not json_string:
            messagebox.showerror("Error", "Please provide JSON resume data")
            return
        
        if not output_filename:
            messagebox.showerror("Error", "Please specify an output filename")
            return
        
        try:
            # Validate JSON
            json.loads(json_string)
            
            # Generate resume
            success = self.generate_resume_from_json(json_string, output_filename)
            
            if success:
                messagebox.showinfo("Success", "Resume generated successfully!")
                self.status_var.set(f"Resume saved to: {output_filename}")
            else:
                messagebox.showerror("Error", "Failed to generate resume")
        except json.JSONDecodeError:
            messagebox.showerror("Error", "Invalid JSON format")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
    
    def generate_resume_from_json(self, json_string, output_filename):
        try:
            data = json.loads(json_string)
            doc = Document()

            # === Styling & Layout ===
            section = doc.sections[0]
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)

            def add_centered_paragraph(text, bold=False, size=12):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(size)
                p.paragraph_format.space_after = Pt(2)

            def add_section_heading(text):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                p_border = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                p_border.append(bottom)
                p._p.get_or_add_pPr().append(p_border)
                p.paragraph_format.space_after = Pt(4)

            def add_bullet_points(items):
                for item in items:
                    p = doc.add_paragraph(style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.left_indent = Inches(0.25)
                    run = p.add_run(item)
                    run.font.size = Pt(11)

            def add_hyperlinked_paragraph(doc, text_parts):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(4)
                for idx, (display_text, url) in enumerate(text_parts):
                    r_id = doc.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    new_run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), '0000FF')
                    rPr.append(color)
                    underline = OxmlElement('w:u')
                    underline.set(qn('w:val'), 'single')
                    rPr.append(underline)
                    new_run.append(rPr)
                    text = OxmlElement('w:t')
                    text.text = display_text
                    new_run.append(text)
                    hyperlink.append(new_run)
                    paragraph._p.append(hyperlink)
                    if idx != len(text_parts) - 1:
                        paragraph.add_run(" | ")

            # === HEADER ===
            add_centered_paragraph(data['name'], bold=True, size=14)
            add_centered_paragraph(data.get('title', ''), size=11)

            contact = data.get('contact', {})
            contact_parts = []
            if contact.get('portfolio'):
                contact_parts.append(('Portfolio', contact['portfolio']))
            if contact.get('linkedin'):
                contact_parts.append(('LinkedIn', contact['linkedin']))
            if contact.get('email'):
                contact_parts.append((contact['email'], f"mailto:{contact['email']}"))
            if contact.get('phone'):
                contact_parts.append((contact['phone'], f"tel:{contact['phone']}"))

            if data.get('portfolio'):
                contact_parts.append(('Portfolio', data['portfolio']))
            if data.get('linkedin'):
                contact_parts.append(('LinkedIn', data['linkedin']))
            if data.get('email') and ('email' not in [part[0] for part in contact_parts]):
                contact_parts.append((data['email'], f"mailto:{data['email']}"))
            if data.get('phone') and ('phone' not in [part[0] for part in contact_parts]):
                contact_parts.append((data['phone'], f"tel:{data['phone']}"))

            if contact_parts:
                add_hyperlinked_paragraph(doc, contact_parts)

            # === PROFESSIONAL SUMMARY ===
            if data.get('professional_summary'):
                add_section_heading("Professional Summary")
                add_bullet_points(data['professional_summary'])

            # === TECHNICAL SKILLS ===
            if data.get('technical_skills'):
                add_section_heading("Technical Skills")
                for category, skills in data['technical_skills'].items():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f"• {category}: ")
                    run.bold = True
                    p.add_run(", ".join(skills))

            # === EXPERIENCE ===
            if data.get('experience'):
                add_section_heading("Professional Experience")
                for job in data['experience']:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Role: {job['role']}")
                    run.bold = True
                    run.font.size = Pt(11)
                    p.paragraph_format.space_after = Pt(0)

                    p = doc.add_paragraph()
                    p.paragraph_format.tab_stops.clear_all()
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3))
                    run_left = p.add_run(f"Client: {job['company']}")
                    run_left.bold = True
                    run_left.font.size = Pt(11)
                    if job.get('duration'):
                        run_right = p.add_run(f"\t{job['duration']}")
                        run_right.bold = True
                        run_right.font.size = Pt(10)
                        run_right.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.space_after = Pt(4)

                    if job.get('project_overview'):
                        p = doc.add_paragraph()
                        run = p.add_run("Project Overview: ")
                        run.bold = True
                        p.add_run(job['project_overview'])
                        p.paragraph_format.space_after = Pt(4)

                    if job.get('responsibilities'):
                        p = doc.add_paragraph()
                        run = p.add_run("Responsibilities: ")
                        run.bold = True
                        p.paragraph_format.space_after = Pt(2)
                        add_bullet_points(job['responsibilities'])

                    if job.get('environment'):
                        p = doc.add_paragraph()
                        run = p.add_run("Environment: ")
                        run.bold = True
                        p.add_run(", ".join(job['environment']))
                        p.paragraph_format.space_after = Pt(8)

            # === EDUCATION ===
            if data.get('education') and isinstance(data['education'], dict):
                add_section_heading("Education")
                edu = data['education']
                p = doc.add_paragraph()
                line_parts = []
                if edu.get('degree'):
                    line_parts.append(edu['degree'])
                if edu.get('field'):
                    line_parts.append(edu['field'])
                if edu.get('institution'):
                    line_parts.append(f"at {edu['institution']}")
                if edu.get('year'):
                    line_parts.append(f"({edu['year']})")
                if line_parts:
                    p.add_run(", ".join(line_parts))
                p.paragraph_format.space_after = Pt(2)

            # === CERTIFICATIONS ===
            if data.get('certifications'):
                add_section_heading("Certifications")
                for cert in data['certifications']:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(cert)
                    run.font.size = Pt(11)
                    p.paragraph_format.space_after = Pt(2)

            # === Save DOCX ===
            doc.save(output_filename)
            print(f"DOCX saved to: {output_filename}")

            # === Convert DOCX to PDF ===
            pdf_filename = output_filename.replace('.docx', '.pdf')
            try:
                subprocess.run([
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(output_filename),
                    output_filename
                ], check=True)
            except Exception as e:
                print(f"⚠️ Could not convert to PDF: {e}")
                return False

            # Create directory structure if it doesn't exist
            today = datetime.now().strftime("%Y-%m-%d")
            local_path = os.path.join(os.path.dirname(output_filename), today)
            os.makedirs(local_path, exist_ok=True)

            # Move files to final location
            final_docx_path = os.path.join(local_path, os.path.basename(output_filename))
            final_pdf_path = os.path.join(local_path, os.path.basename(pdf_filename))
            
            shutil.move(output_filename, final_docx_path)
            shutil.move(pdf_filename, final_pdf_path)
            
            print(f"✅ Files successfully saved to:")
            print(f"DOCX: {final_docx_path}")
            print(f"PDF: {final_pdf_path}")
            
            return True
            
        except Exception as e:
            print(f"❌ Error generating resume: {e}")
            return False

if __name__ == "__main__":
    root = tk.Tk()
    app = ResumeGeneratorApp(root)
    root.mainloop()